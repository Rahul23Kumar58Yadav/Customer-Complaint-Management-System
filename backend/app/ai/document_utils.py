"""
Lightweight document text extraction.
"""
import email
import io
import logging
from email import policy
from dataclasses import dataclass, field

from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    text: str
    page_count: int | None = None
    warnings: list[str] = field(default_factory=list)


def extract_text_from_bytes(filename: str, content: bytes) -> str:
    return extract_with_metadata(filename, content).text


def extract_with_metadata(filename: str, content: bytes) -> ExtractionResult:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf(content)
    elif ext == "docx":
        return _extract_docx(content)
    elif ext == "eml":
        return _extract_eml(content)
    elif ext == "txt":
        return _extract_txt(content)
    else:
        logger.warning("Unrecognized extension '%s' for file '%s', attempting best-effort decode", ext, filename)
        return ExtractionResult(text=_decode_best_effort(content), warnings=["unrecognized_extension"])


def _decode_best_effort(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")


def _extract_pdf(content: bytes) -> ExtractionResult:
    warnings = []
    try:
        reader = PdfReader(io.BytesIO(content))
    except PdfReadError as exc:
        logger.error("Failed to open PDF: %s", exc)
        return ExtractionResult(text="", page_count=0, warnings=[f"pdf_open_failed: {exc}"])

    if reader.is_encrypted:
        try:
            reader.decrypt("")
            warnings.append("pdf_was_encrypted_unlocked_with_empty_password")
        except Exception:  # noqa: BLE001
            logger.warning("PDF is encrypted and could not be auto-unlocked")
            return ExtractionResult(
                text="", page_count=len(reader.pages), warnings=["pdf_encrypted_could_not_unlock"]
            )

    pages_text = []
    for i, page in enumerate(reader.pages):
        try:
            pages_text.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            logger.warning("Failed to extract text from PDF page %d: %s", i, exc)
            warnings.append(f"page_{i}_extraction_failed")

    return ExtractionResult(text="\n".join(pages_text), page_count=len(reader.pages), warnings=warnings)


def _extract_docx(content: bytes) -> ExtractionResult:
    try:
        doc = DocxDocument(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        logger.error("Failed to open DOCX: %s", exc)
        return ExtractionResult(text="", warnings=[f"docx_open_failed: {exc}"])

    paragraphs = [p.text for p in doc.paragraphs]
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            table_text.append(" | ".join(cell.text for cell in row.cells))

    text = "\n".join(paragraphs + table_text)
    return ExtractionResult(text=text, warnings=["docx_tables_included"] if table_text else [])


def _extract_eml(content: bytes) -> ExtractionResult:
    try:
        msg = email.message_from_bytes(content, policy=policy.default)
    except Exception as exc:  # noqa: BLE001
        logger.error("Failed to parse EML: %s", exc)
        return ExtractionResult(text="", warnings=[f"eml_parse_failed: {exc}"])

    parts = [f"From: {msg.get('from', '')}", f"Subject: {msg.get('subject', '')}"]
    warnings = []
    body = msg.get_body(preferencelist=("plain", "html"))
    if body:
        parts.append(body.get_content())
    else:
        warnings.append("eml_no_body_found")

    attachments = [part.get_filename() for part in msg.iter_attachments() if part.get_filename()]
    if attachments:
        parts.append(f"[Attachments not parsed: {', '.join(attachments)}]")
        warnings.append("eml_attachments_not_parsed")

    return ExtractionResult(text="\n".join(parts), warnings=warnings)


def _extract_txt(content: bytes) -> ExtractionResult:
    return ExtractionResult(text=_decode_best_effort(content))